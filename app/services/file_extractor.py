"""
File extraction service for PDF, DOCX, and TXT files.
"""

import os
from dataclasses import dataclass
from typing import List, Optional
import mimetypes


@dataclass
class ExtractedTextSection:
    """Extracted text with optional source page metadata."""

    text: str
    page_no: Optional[int] = None


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from a TXT file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file."""
    return "\n".join(section.text for section in extract_text_sections_from_pdf(file_path))


def extract_text_sections_from_pdf(file_path: str) -> List[ExtractedTextSection]:
    """Extract text from a PDF file, preserving one-based page numbers."""
    try:
        import PyPDF2
    except ImportError:
        raise ImportError("PyPDF2 is required to extract text from PDF files")
    
    sections = []
    try:
        with open(file_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page_num, page in enumerate(pdf_reader.pages, start=1):
                sections.append(
                    ExtractedTextSection(
                        text=page.extract_text() or "",
                        page_no=page_num,
                    )
                )
    except Exception as e:
        raise ValueError(f"Error extracting text from PDF: {str(e)}")
    
    return sections


def get_page_count(file_path: str, file_type: str) -> Optional[int]:
    """Return the physical page count when the file format exposes it."""
    file_type = file_type.lower().strip('.')

    if file_type != "pdf":
        return None

    try:
        import PyPDF2
    except ImportError:
        raise ImportError("PyPDF2 is required to count PDF pages")

    try:
        with open(file_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            return len(pdf_reader.pages)
    except Exception as e:
        raise ValueError(f"Error counting PDF pages: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required to extract text from DOCX files")
    
    text = []
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                if row_text:
                    text.append(" | ".join(row_text))
    except Exception as e:
        raise ValueError(f"Error extracting text from DOCX: {str(e)}")
    
    return "\n".join(text)


def extract_text(file_path: str, file_type: str) -> str:
    """
    Extract text from a file based on its type.
    
    Args:
        file_path: Path to the file
        file_type: Type of file (pdf, docx, txt)
    
    Returns:
        Extracted text content
    
    Raises:
        ValueError: If file type is not supported or extraction fails
    """
    file_type = file_type.lower().strip('.')
    
    return "\n".join(section.text for section in extract_text_sections(file_path, file_type))


def extract_text_sections(file_path: str, file_type: str) -> List[ExtractedTextSection]:
    """
    Extract text sections from a file, preserving page metadata when available.
    """
    file_type = file_type.lower().strip('.')

    if file_type == "txt":
        return [ExtractedTextSection(text=extract_text_from_txt(file_path))]
    elif file_type == "pdf":
        return extract_text_sections_from_pdf(file_path)
    elif file_type == "docx":
        return [ExtractedTextSection(text=extract_text_from_docx(file_path))]
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def get_file_type(filename: str) -> Optional[str]:
    """Get file type from filename."""
    _, ext = os.path.splitext(filename)
    return ext.lstrip('.').lower() if ext else None


def validate_file(filename: str, max_size: int, allowed_types: list) -> tuple[bool, str]:
    """
    Validate a file before processing.
    
    Args:
        filename: Name of the file
        max_size: Maximum file size in bytes
        allowed_types: List of allowed file types (e.g., ['pdf', 'docx', 'txt'])
    
    Returns:
        Tuple of (is_valid, message)
    """
    file_type = get_file_type(filename)
    
    if not file_type:
        return False, "Could not determine file type"
    
    if file_type not in allowed_types:
        return False, f"File type '{file_type}' not allowed. Allowed types: {', '.join(allowed_types)}"
    
    return True, "Valid"
